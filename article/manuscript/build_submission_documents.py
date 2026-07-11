#!/usr/bin/env python3
"""Build the revised Earth Science Informatics Word submission package.

The manuscript source is kept in Markdown for reviewability. This builder adds
Word-native tables, embeds regenerated figures with alt text, applies inline
monospaced character formatting to source-code terms, and creates portrait
documents with page-number fields. It deliberately carries unresolved
author-controlled declarations and DOI status as unresolved rather than
inventing them.
"""

from __future__ import annotations

import csv
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PACKAGE_ROOT = Path(__file__).resolve().parents[2]
WORKSPACE_ROOT = PACKAGE_ROOT.parents[1]
ARTICLE_ROOT = PACKAGE_ROOT / "article"
SUBMISSION_ROOT = WORKSPACE_ROOT / "blueterra_ESI_submission"
FIGURE_DIR = ARTICLE_ROOT / "figures" / "output"
TABLE_DIR = ARTICLE_ROOT / "tables" / "results"
REFERENCE_FILE = ARTICLE_ROOT / "references" / "references.csv"
SOURCE_MANUSCRIPT = ARTICLE_ROOT / "manuscript" / "blueterra_ESI_manuscript.md"

TITLE = "blueterra: an R workflow for geomorphometric analysis of submerged terrain"
AUTHOR = "Elvin Cordero"
AFFILIATIONS = [
    "Department of Marine Sciences, University of Puerto Rico at Mayagüez, Mayagüez, Puerto Rico, USA",
    "SeaMount Geospatial Labs, Brooklyn, New York, USA",
]
EMAIL = "elvin.cordero1@upr.edu"

TABLE_SPECS = {
    1: {
        "file": "table1_core_metric_definitions.csv",
        "caption": "Core metric definitions and implementation boundaries.",
        "note": "Operational definitions describe version 0.2.0. The labels do not imply direct measurement of an environmental process.",
        "headers": ["metric", "implementation source", "operational definition", "units", "neighborhood or scale", "principal interpretation constraint"],
        "keys": ["metric", "implementation_source", "operational_definition", "units", "neighbourhood_or_scale", "principal_interpretation_constraint"],
        "widths": [0.75, 0.96, 1.45, 0.65, 1.12, 1.37],
    },
    2: {
        "file": "table2_verification_summary.csv",
        "caption": "Numerical, upstream-wrapper, and functional verification summary.",
        "note": "Direct terra comparisons are upstream-wrapper equivalence, not independent cross-implementation validation. Detailed records are in Online Resource 2.",
        "headers": ["metric family", "reference behavior or comparator", "evaluated domain", "edge policy", "maximum error", "tolerance", "result"],
        "keys": ["metric_family", "reference_behavior_or_comparator", "evaluated_domain", "edge_policy", "maximum_error", "tolerance", "result"],
        "widths": [0.82, 1.20, 0.92, 1.20, 0.62, 0.60, 0.44],
    },
    3: {
        "file": "table3_sensitivity_computational_summary.csv",
        "caption": "Selected sensitivity and computational results.",
        "note": "Detailed sensitivity, timing, environment, and clean-environment records are supplied in Online Resource 2. Benchmark hardware is not a minimum requirement.",
        "headers": ["assessment", "scenario", "setting", "observed result", "interpretation"],
        "keys": ["assessment", "scenario", "setting", "observed_result", "interpretation"],
        "widths": [0.67, 1.20, 1.48, 1.18, 1.47],
    },
}

FIGURES = {
    1: {
        "file": "Fig1_architecture_workflow.png",
        "caption": "Architecture of the version 0.2.0 workflow. An input raster is prepared before a metric stack is derived from terra-wrapped operations and local formulas. The metric catalog accompanies the stack as metadata, while aligned custom or external layers may join it before spatial summaries, transects, isobath corridors, tables, maps, and profiles are produced.",
        "alt": "Workflow diagram from input raster through preparation and aligned metric stack to spatial analyses and outputs, with metadata catalog and optional custom layers.",
    },
    2: {
        "file": "Fig2_end_to_end_terrain_workflow.png",
        "caption": "Documented NOAA BlueTopo tile BH54S4ZB example in southwest Puerto Rico: elevation with a Puerto Rico location inset and 500 m scale bar, slope, fine-scale BPI, VRM-style rugosity, four-neighbor Laplacian-style index, and slope-based surface-area ratio. All maps use EPSG:6348 and 4 m grid spacing; stored values are negative elevations in meters, no smoothing or vertical transformation was applied, and BPI/VRM use 3 by 3 cells.",
        "alt": "Six terrain maps from a documented 4 meter BlueTopo elevation crop, including location inset and scale bar.",
    },
    3: {
        "file": "Fig3_catalog_and_summaries.png",
        "caption": "Metric metadata and structured summaries for the documented stack. The process-group strip assigns the five derived layers to transparent descriptor groups; it does not assert derivation of catalog-only external layers. The lower panels show author-created analysis windows and distributions of slope values across monotonically ordered elevation bands, with contributing cell counts and median markers.",
        "alt": "Process-group strip, two analysis windows on elevation map, and slope box plots by ordered elevation band.",
    },
    4: {
        "file": "Fig4_transects_cross_sections.png",
        "caption": "Terrain-oriented transects for the author-created slope window. Six lines were clipped to the polygon at 250 m spacing and share a slope-weighted surface-derived angle of 84.8 degrees (resultant length 0.771; 103,666 contributing cells). Matching colors identify the map lines and profiles; depth is displayed downward and distance begins at the shallower endpoint.",
        "alt": "Elevation map with a white source polygon and six colored transects, paired with matching colored depth profiles that deepen from left to right.",
    },
    5: {
        "file": "Fig5_isobath_corridors.png",
        "caption": "Independent corridors around −100, −300, and −600 m source contours. Black dashed lines are source isobaths and orange lines are corridor boundaries. Each uses a 20 m one-sided buffer (40 m nominal full width, ten 4 m cells); no overlap was detected for these selected corridors. Box plots show slope distributions, with corridor area and contributing-cell count above each group.",
        "alt": "Elevation map with three dashed black contour lines and orange corridor boundaries, paired with slope box plots for three depth contours.",
    },
    6: {
        "file": "Fig6_sensitivity.png",
        "caption": "Separated sensitivity results. Comparable BPI maps use identical symmetric color limits for unsmoothed 3 by 3 BPI, 3 by 3 mean-smoothed BPI, and unsmoothed 11 by 11 BPI. Lower panels report median absolute cellwise differences and Spearman correlations in separate metric-specific facets for grid, preprocessing, and focal-neighborhood comparisons.",
        "alt": "Three BPI maps with identical diverging color scale above metric-specific sensitivity and correlation facets.",
    },
    7: {
        "file": "Fig7_validation_and_agreement.png",
        "caption": "Controlled validation and implementation agreement. Synthetic planar, convex, and center-relief surfaces illustrate tested reference behavior. Observed values equal expected values for selected analytical cases, and the maximum error-to-tolerance ratios remain below the threshold of one on a linear scale; exact zero differences are retained at zero.",
        "alt": "Synthetic planar, convex Laplacian, and BPI surfaces above observed-versus-expected values and zero error-to-tolerance ratios below a threshold of one.",
    },
}

CITED_KEYS = [
    "Baston2025_exactextractr",
    "DeReuEtAl2013",
    "DolanLucieer2014",
    "ErdeyHeydorn2008",
    "Hengl2006",
    "HijmansBrownBarbosa2026_terra",
    "Horn1981",
    "IHO_S44_2022",
    "IlichMisiukLecoursMurawski2023",
    "Jenness2004",
    "LecoursDevillersLucieerBrown2017",
    "LecoursDolanMicallefLucieer2016",
    "Lindsay2016",
    "LundbladEtAl2006",
    "MarwickBoettigerMullen2018",
    "MisiukBrown2024",
    "MisiukLecoursBell2018",
    "MisiukLecoursDolanRobert2021",
    "NOAA_BlueTopo_BH54S4ZB_20251117",
    "NOAA_BlueTopo_Product",
    "Pebesma2018",
    "RileyDeGloriaElliot1999",
    "SappingtonLongshoreThompson2007",
    "ShermanEtAl2010",
    "ThompsonBellButler2001",
    "WalbridgeEtAl2018",
    "WilsonEtAl2007",
    "WuBrown2025_whitebox",
]


def set_font(run, size=10, bold=None, italic=None, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def clear_paragraph(paragraph):
    for child in list(paragraph._element):
        paragraph._element.remove(child)


def set_section_geometry(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.36)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(paragraph)


def configure_document(doc):
    set_section_geometry(doc.sections[0])
    add_footer(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.0
    for name, size, before, after in [("Heading 1", 12, 10, 4), ("Heading 2", 11, 8, 3), ("Heading 3", 10, 6, 2)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Reference" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
    if "Inline Code" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Inline Code", WD_STYLE_TYPE.CHARACTER)
        style.font.name = "Courier New"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        style.font.size = Pt(8.5)


def add_inline_text(paragraph, text, size=10, reference=False):
    """Render backtick spans as a real monospaced character style."""
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.style = "Inline Code"
            set_font(run, 8.5 if not reference else 8, name="Courier New")
        else:
            run = paragraph.add_run(part)
            set_font(run, 8.5 if reference else size)


def add_text_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT, reference=False):
    paragraph = doc.add_paragraph(style="Reference" if reference else "Normal")
    paragraph.alignment = align
    paragraph.paragraph_format.widow_control = True
    add_inline_text(paragraph, text, size=10, reference=reference)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_font(run, {1: 12, 2: 11, 3: 10}[min(level, 3)], bold=True)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=36, start=45, bottom=36, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table, widths):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = int(sum(widths) * 1440)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(int(width * 1440)))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_caption(doc, prefix, number, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    lead = paragraph.add_run(f"{prefix} {number}. ")
    set_font(lead, 9, bold=True)
    add_inline_text(paragraph, text, size=9)
    return paragraph


def display_value(value):
    text = str(value).strip()
    if text.lower() in {"na", "nan", "none"}:
        return ""
    return text


def table_rows(number):
    spec = TABLE_SPECS[number]
    with (TABLE_DIR / spec["file"]).open(newline="", encoding="utf-8") as handle:
        records = list(csv.DictReader(handle))
    return [[record.get(key, "") for key in spec["keys"]] for record in records]


def add_table(doc, number):
    spec = TABLE_SPECS[number]
    add_caption(doc, "Table", number, spec["caption"])
    values = table_rows(number)
    table = doc.add_table(rows=1, cols=len(spec["headers"]))
    table.style = "Table Grid"
    table.autofit = False
    set_table_fixed_layout(table, spec["widths"])
    header = table.rows[0]
    for cell, label, width in zip(header.cells, spec["headers"], spec["widths"]):
        set_cell_width(cell, width)
        set_cell_margins(cell)
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 0.88
        add_inline_text(paragraph, label, size=8)
        for run in paragraph.runs:
            run.bold = True
    set_repeat_table_header(header)
    for row_values in values:
        row = table.add_row()
        prevent_row_split(row)
        for cell, value, width in zip(row.cells, row_values, spec["widths"]):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.88
            add_inline_text(paragraph, display_value(value), size=8)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(3)
    note.paragraph_format.space_after = Pt(5)
    run = note.add_run("Note. ")
    set_font(run, 8, italic=True)
    add_inline_text(note, spec["note"], size=8)
    for run in note.runs[1:]:
        run.italic = True


def set_image_alt_text(shape, alt_text):
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text[:120])


def add_figure(doc, number):
    spec = FIGURES[number]
    image = FIGURE_DIR / spec["file"]
    if not image.exists():
        raise FileNotFoundError(image)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(image), width=Inches(6.55))
    set_image_alt_text(shape, spec["alt"])
    add_caption(doc, "Fig.", number, spec["caption"])


def add_author_with_superscripts(paragraph):
    run = paragraph.add_run(AUTHOR)
    set_font(run, 11)
    superscript = paragraph.add_run("1,2")
    set_font(superscript, 8)
    superscript.font.superscript = True


def add_affiliation(paragraph, index, text):
    index_run = paragraph.add_run(str(index))
    set_font(index_run, 8)
    index_run.font.superscript = True
    text_run = paragraph.add_run(f" {text}")
    set_font(text_run, 10)


def add_manuscript_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run(TITLE)
    set_font(run, 16, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    add_author_with_superscripts(paragraph)
    for index, affiliation in enumerate(AFFILIATIONS, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        add_affiliation(paragraph, index, affiliation)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    run = paragraph.add_run(f"Correspondence: {EMAIL}")
    set_font(run, 10)
    doc.add_page_break()


def load_references():
    with REFERENCE_FILE.open(newline="", encoding="utf-8") as handle:
        records = {record["citation_key"]: record["bibliography_entry"] for record in csv.DictReader(handle)}
    missing = [key for key in CITED_KEYS if key not in records]
    if missing:
        raise ValueError(f"Missing verified bibliography records: {', '.join(missing)}")
    return records


def add_references(doc):
    records = load_references()
    for key in CITED_KEYS:
        add_text_paragraph(doc, records[key], reference=True)


def parse_manuscript(doc):
    heading_re = re.compile(r"^(#{1,3})\s+(.*)$")
    buffer = []
    skip_reference_source = False

    def flush():
        nonlocal buffer
        if buffer:
            add_text_paragraph(doc, " ".join(buffer).strip())
            buffer = []

    for raw_line in SOURCE_MANUSCRIPT.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        marker = re.fullmatch(r"\[(FIGURE|TABLE)\s+(\d+)\]", line)
        heading = heading_re.match(line)
        if heading:
            # A heading after the manuscript's reference-source block resumes
            # ordinary source parsing (for example, Software Files).
            if skip_reference_source:
                skip_reference_source = False
            flush()
            text = heading.group(2)
            add_heading(doc, text, len(heading.group(1)))
            if text == "References":
                add_references(doc)
                skip_reference_source = True
            continue
        if skip_reference_source:
            continue
        if marker:
            flush()
            kind, number = marker.group(1), int(marker.group(2))
            if kind == "FIGURE":
                add_figure(doc, number)
            else:
                add_table(doc, number)
            continue
        if not line:
            flush()
            continue
        buffer.append(line)
    flush()


def build_manuscript():
    doc = Document()
    configure_document(doc)
    add_manuscript_title_page(doc)
    parse_manuscript(doc)
    path = SUBMISSION_ROOT / "blueterra_ESI_manuscript.docx"
    doc.save(path)
    return path


def build_title_page():
    doc = Document()
    configure_document(doc)
    for _ in range(3):
        doc.add_paragraph()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run(TITLE)
    set_font(run, 16, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_author_with_superscripts(paragraph)
    for index, affiliation in enumerate(AFFILIATIONS, start=1):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(2)
        add_affiliation(paragraph, index, affiliation)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    run = paragraph.add_run(f"Corresponding author: Elvin Cordero ({EMAIL})")
    set_font(run, 10)
    add_heading(doc, "Acknowledgments", 1)
    add_text_paragraph(doc, "The author acknowledges the Department of Marine Sciences at the University of Puerto Rico at Mayagüez for logistical and academic support during graduate research and dissertation-related activities. The author also thanks SeaMount Geospatial Labs for providing access to computational resources and infrastructure that supported this project.")
    add_heading(doc, "Author contributions", 1)
    add_text_paragraph(doc, "An author-contributions statement has not been supplied by the author and remains unresolved.")
    path = SUBMISSION_ROOT / "blueterra_ESI_title_page.docx"
    doc.save(path)
    return path


def build_cover_letter():
    doc = Document()
    configure_document(doc)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(14)
    run = paragraph.add_run("11 July 2026")
    set_font(run, 10)
    add_text_paragraph(doc, "Editor-in-Chief\nEarth Science Informatics")
    add_text_paragraph(doc, "Dear Editor,")
    add_text_paragraph(doc, f"Please consider the enclosed Software Article, “{TITLE},” for publication in Earth Science Informatics.")
    add_text_paragraph(doc, "The article presents blueterra as an R workflow for transparent preparation and geomorphometric description of submerged terrain. It combines a bounded metric stack with metric metadata, spatial summaries, terrain-oriented transects, isobath corridors, visualization, and aligned user-defined layers. The evaluation reports controlled numerical checks, direct upstream-wrapper equivalence, functional spatial checks, documented BlueTopo example analyses, scale sensitivity, and computational measurements.")
    add_text_paragraph(doc, "The software is free under the MIT License. Version 0.2.0 is identified by release tag v0.2.0, and the associated source, data-preparation materials, scripts, results, and environment records are included with the article materials. A persistent archive DOI remains pending and is transparently recorded in the compliance checklist.")
    add_text_paragraph(doc, "The manuscript is limited to demonstrated software behavior and does not claim physical-process prediction, ecological inference, or a performance advantage over other software. It is offered because its documented implementation boundary, reproducible spatial workflow, and evaluation align with the journal’s Earth science informatics scope.")
    add_text_paragraph(doc, "Sincerely,")
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Elvin Cordero")
    set_font(run, 10)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(EMAIL)
    set_font(run, 10)
    path = SUBMISSION_ROOT / "blueterra_ESI_cover_letter.docx"
    doc.save(path)
    return path


def reset_directory(path):
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def ignore_superseded_article_materials(directory, names):
    """Exclude retained pre-revision evidence from the submission archive copy."""
    relative = Path(directory).resolve().relative_to(ARTICLE_ROOT)
    ignored = set()
    if relative == Path("reproducibility") and "results" in names:
        # Keep the result directory itself: its tag-based child is selected below.
        pass
    if relative == Path("reproducibility") / "results":
        ignored.update(name for name in names if "_head-4afb4b58" in name)
    if relative == Path("references"):
        ignored.update({
            "assemble_reference_audit.py", "reference_audit_foundational.csv",
            "reference_audit_marine.csv", "reference_audit_revision.csv",
            "reference_audit_software.csv",
        })
    if relative == Path("tables"):
        ignored.add("collect_tables_4_and_5.R")
    if relative == Path("tables") / "results":
        ignored.update({
            "table1_core_metric_implementation.csv", "table2_source_records.csv",
            "table2_validation_wrapper_agreement.csv", "table3_clean_environment_reproducibility.csv",
            "table3_source_comparison.csv",
        })
    return ignored


def copy_submission_components():
    for name in ["figures", "tables", "scripts", "environment", "supplementary"]:
        reset_directory(SUBMISSION_ROOT / name)
    figures_out = SUBMISSION_ROOT / "figures"
    tables_out = SUBMISSION_ROOT / "tables"
    scripts_out = SUBMISSION_ROOT / "scripts"
    environment_out = SUBMISSION_ROOT / "environment"
    supplementary_out = SUBMISSION_ROOT / "supplementary"
    for source in FIGURE_DIR.glob("Fig*.*"):
        if source.suffix.lower() in {".pdf", ".png", ".tiff"}:
            shutil.copy2(source, figures_out / source.name)
    for filename in [
        "table1_core_metric_definitions.csv", "table2_verification_summary.csv",
        "table3_sensitivity_computational_summary.csv", "table2_numeric_source_records.csv",
        "table2_functional_source_records.csv", "table3_full_sensitivity_records.csv",
        "table3_full_benchmark_records.csv",
    ]:
        shutil.copy2(TABLE_DIR / filename, tables_out / filename)
    for filename in ["reference_audit.csv", "references.csv"]:
        shutil.copy2(ARTICLE_ROOT / "references" / filename, SUBMISSION_ROOT / filename)
    for filename, target in [
        ("submission_compliance_checklist.md", "submission_compliance_checklist.md"),
        ("provenance_and_change_status.md", "data_provenance_status.md"),
        ("package_modifications.csv", "package_modifications.csv"),
    ]:
        shutil.copy2(ARTICLE_ROOT / filename, SUBMISSION_ROOT / target)
    captions = ARTICLE_ROOT / "supplementary" / "Online_Resource_Captions.md"
    if captions.exists():
        shutil.copy2(captions, supplementary_out / captions.name)
    for relative in ["validation", "reproducibility", "sensitivity", "benchmark", "figures", "tables", "environment", "references", "manuscript", "supplementary", "data_provenance"]:
        source = ARTICLE_ROOT / relative
        destination = scripts_out / relative
        shutil.copytree(
            source,
            destination,
            ignore=lambda directory, names: set(
                shutil.ignore_patterns("*.tiff", "*.png", "*.pdf", "__pycache__", "*.pyc")(directory, names)
            ) | ignore_superseded_article_materials(directory, names),
        )
    environment_results = ARTICLE_ROOT / "environment" / "results"
    if environment_results.exists():
        for source in environment_results.iterdir():
            if source.is_file():
                shutil.copy2(source, environment_out / source.name)
    return supplementary_out


def main():
    SUBMISSION_ROOT.mkdir(parents=True, exist_ok=True)
    manuscript = build_manuscript()
    title_page = build_title_page()
    cover_letter = build_cover_letter()
    supplementary = copy_submission_components()
    print(manuscript)
    print(title_page)
    print(cover_letter)
    print(supplementary)


if __name__ == "__main__":
    main()
